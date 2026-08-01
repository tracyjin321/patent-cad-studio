"""Safe in-memory text extraction for supported patent documents."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_DOCUMENT_SUFFIXES = {".pdf", ".docx", ".txt"}


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\x00", "").splitlines()]
    return "\n".join(line for line in lines if line).strip()


def _extract_pdf(content: bytes) -> str:
    if not content.startswith(b"%PDF"):
        raise ValueError("PDF 文件格式无效")
    reader = PdfReader(BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        blocks.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
    return "\n".join(blocks)


def _extract_txt(content: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("TXT 文件编码无法识别，请使用 UTF-8 或 GB18030")


def extract_document_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_DOCUMENT_SUFFIXES:
        raise ValueError("仅支持 PDF、DOCX 和 TXT 文档")
    try:
        if suffix == ".pdf":
            text = _extract_pdf(content)
        elif suffix == ".docx":
            text = _extract_docx(content)
        else:
            text = _extract_txt(content)
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"{suffix[1:].upper()} 文档解析失败") from exc
    normalized = _normalize_text(text)
    if not normalized:
        raise ValueError("文档中未提取到可用文字；扫描版 PDF 暂不支持 OCR")
    return normalized
